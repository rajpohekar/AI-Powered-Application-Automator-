import docx

def create_resume():
    doc = docx.Document()
    
    doc.add_heading('Jane Doe', 0)
    
    doc.add_paragraph('Email: janedoe@example.com | Phone: (555) 123-4567 | Location: San Francisco, CA')
    doc.add_paragraph('GitHub: github.com/janedoe | LinkedIn: linkedin.com/in/janedoe')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(
        'Passionate Senior Software Engineer with 6 years of experience building scalable web applications. '
        'Expert in Node.js, Python, React, and cloud architectures. Looking for roles with a target compensation of $150,000.'
    )
    
    doc.add_heading('Experience', level=1)
    p1 = doc.add_paragraph('Senior Software Engineer - TechCorp (2022 - Present)')
    p1.bold = True
    doc.add_paragraph(
        '- Led a team of 4 developers to redesign a high-traffic microservice gateway in Node.js.\n'
        '- Improved database query performances by 40% using Redis caching and Mongo optimizations.\n'
        '- Built RAG pipelines integrating Qdrant vector databases and OpenAI endpoints.'
    )
    
    doc.add_heading('Education', level=1)
    p2 = doc.add_paragraph('B.S. in Computer Science - Stanford University (2016 - 2020)')
    p2.bold = True
    
    doc.save('candidate_resume.docx')
    print("Mock resume candidate_resume.docx created successfully!")

if __name__ == '__main__':
    create_resume()
